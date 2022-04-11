import time
import docx
import os
import datetime
import pandas as pd
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from PIL import Image 
from docx.shared import Inches #设置图像大小
from selenium.webdriver.common.action_chains import ActionChains

chrome_options = Options()
chrome_options.add_argument('--start-maximized')
chrome_options.add_experimental_option("debuggerAddress", "127.0.0.1:9222")
driver = webdriver.Chrome(chrome_options=chrome_options)
print(driver.title)

# 屏幕缩放比例
pingmu_suofang = 1.5

# YYYYMMDD
current_time = datetime.datetime.now()
sub_day = datetime.timedelta(days=1)
sub_time = (current_time - sub_day).strftime("%Y%m%d")
sub_time_gan = (current_time - sub_day).strftime("%Y-%m-%d")
ssub_time_gan = (current_time - sub_day - sub_day).strftime("%Y-%m-%d")

def quyu_screenshot(elem4, pic1, pic2, is_two=True):
    """
    截屏：两个图片
    """
    locations = elem4.location
    print(locations)
    #图片大小
    sizes = elem4.size
    print(sizes)
    # 构造指数的位置
    if is_two:
        # rangle = (int(locations['x']), int(locations['y']) - 20, int(locations['x'] + 2 * sizes['width']), int(locations['y'] + sizes['height']))
        rangle = (int(locations['x']) * pingmu_suofang, int(locations['y']) * pingmu_suofang - 20, int(locations['x'] + 2 * sizes['width']) * pingmu_suofang, int(locations['y'] + sizes['height']) * pingmu_suofang)
    else:
        rangle = (int(locations['x']) * pingmu_suofang, int(locations['y']) * pingmu_suofang - 20, int(locations['x'] + sizes['width']) * pingmu_suofang, int(locations['y'] + sizes['height']) * pingmu_suofang)
    print(rangle)
    save_path = pic1
    driver.save_screenshot(save_path)
    # 打开截图切割
    img = Image.open(save_path)
    jpg = img.convert('RGB')
    jpg = img.crop(rangle)
    path = pic2
    jpg.save(path)
    print("图片截取成功!")

def cal_rate(ribao_file, sheet_nums, col_nums):
    """
    计算升降比
    """
    data_input = pd.read_excel(ribao_file, sheet_name=sheet_nums)

    try:
        cur_jubao_5ren = float(data_input[data_input["日期"] == sub_time_gan][col_nums])
        last_jubao_5ren = float(data_input[data_input["日期"] == ssub_time_gan][col_nums])
        rate_jubao_5ren = cur_jubao_5ren / last_jubao_5ren - 1
    except:
        cur_jubao_5ren = 0
        last_jubao_5ren = 0
        rate_jubao_5ren = 0

    return cur_jubao_5ren, rate_jubao_5ren

def write_content_in_doc(rate_value):
    """计算上涨还是下降"""
    if rate_value > 0:
        write_content1 = "上涨%.1f%s↑" % (abs(rate_value * 100), "%")
    else:
        write_content1 = "下降%.1f%s↓" % (abs(rate_value * 100), "%")
    return write_content1
    

""" 0.查找报表路径"""
# elem1 = driver.find_element_by_id("report_my_nav")
# elem1.click()
# print("click: %s" % elem1.text)
# time.sleep(2)

# # 报表1.0：点击不成功，所以使用js命令
# elem1 = 'document.querySelector("#id_index_sidebar > nav > ul.know_cate_warp.report_all_nav_info > ul:nth-child(11) > li > a").click();'
# driver.execute_script(elem1)
# time.sleep(2)

# # 手游专题
# elem1 = 'document.querySelector("#id_index_sidebar > nav > ul.know_cate_warp.report_all_nav_info > ul:nth-child(11) > li > ul > ul:nth-child(16) > li > a").click();'
# driver.execute_script(elem1)
# time.sleep(2)

# # 王者荣耀日报
# elem1 = 'document.querySelector("#id_index_sidebar > nav > ul.know_cate_warp.report_all_nav_info > ul:nth-child(11) > li > ul > ul:nth-child(16) > li > ul > ul:nth-child(18) > li > a").click();'
# driver.execute_script(elem1)
# time.sleep(2)

# # 日报同步
# elem1 = 'document.querySelector("#id_index_sidebar > nav > ul.know_cate_warp.report_all_nav_info > ul:nth-child(11) > li > ul > ul:nth-child(16) > li > ul > ul:nth-child(18) > li > ul > li > a").click()'
# driver.execute_script(elem1)
# time.sleep(10)

""" 1."""
elem4 = driver.find_element_by_xpath("//*[@id='chart0']/div[1]")
# 滑动滚动条到某个指定的元素
js4 = "arguments[0].scrollIntoView();" 
# 将下拉滑动条滑动到当前div区域
driver.execute_script(js4, elem4)

pic1 = "./result/test_V1.png"
pic2 = "./result/test_V11.png"
quyu_screenshot(elem4, pic1, pic2)

""" 2. """
elem5 = driver.find_element_by_xpath('//*[@id="chart2"]/div[1]')
# 滑动滚动条到某个指定的元素
js4 = "arguments[0].scrollIntoView();" 
# 将下拉滑动条滑动到当前div区域
driver.execute_script(js4, elem5)  

pic1 = "./result/test_V2.png"
pic2 = "./result/test_V22.png"
quyu_screenshot(elem5, pic1, pic2)

elem5 = driver.find_element_by_xpath('//*[@id="chart4"]/div[1]')
js4 = "arguments[0].scrollIntoView();" 
driver.execute_script(js4, elem5)  

pic1 = "./result/test_V2.png"
pic2 = "./result/test_V23.png"
quyu_screenshot(elem5, pic1, pic2, is_two=False)


""" 3. """
elem5 = driver.find_element_by_xpath('//*[@id="chart6"]/div[1]')
# 滑动滚动条到某个指定的元素
js4 = "arguments[0].scrollIntoView();" 
# 将下拉滑动条滑动到当前div区域
driver.execute_script(js4, elem5)  

pic1 = "./result/test_V3.png"
pic2 = "./result/test_V33.png"
quyu_screenshot(elem5, pic1, pic2)

""" 4. """
elem5 = driver.find_element_by_xpath('//*[@id="chart8"]/div[1]')
# 滑动滚动条到某个指定的元素
js4 = "arguments[0].scrollIntoView();" 
# 将下拉滑动条滑动到当前div区域
driver.execute_script(js4, elem5)  

pic1 = "./result/test_V4.png"
pic2 = "./result/test_V44.png"
quyu_screenshot(elem5, pic1, pic2)

""" 5. """
elem5 = driver.find_element_by_xpath('//*[@id="chart12"]/div[1]')
# 滑动滚动条到某个指定的元素
js4 = "arguments[0].scrollIntoView();" 
# 将下拉滑动条滑动到当前div区域
driver.execute_script(js4, elem5)  

pic1 = "./result/test_V5.png"
pic2 = "./result/test_V55.png"
quyu_screenshot(elem5, pic1, pic2)

""" 6. """
elem5 = driver.find_element_by_xpath('//*[@id="chart14"]/div[1]')
# 滑动滚动条到某个指定的元素
js4 = "arguments[0].scrollIntoView();" 
# 将下拉滑动条滑动到当前div区域
driver.execute_script(js4, elem5)

pic1 = "./result/test_V6.png"
pic2 = "./result/test_V66.png"
quyu_screenshot(elem5, pic1, pic2, is_two=False)

""" 7. """
elem5 = driver.find_element_by_xpath('//*[@id="chart15"]/div[1]')
# 滑动滚动条到某个指定的元素
js4 = "arguments[0].scrollIntoView();" 
# 将下拉滑动条滑动到当前div区域
driver.execute_script(js4, elem5)

pic1 = "./result/test_V7.png"
pic2 = "./result/test_V77.png"
quyu_screenshot(elem5, pic1, pic2, is_two=False)

if __name__ == "__main__":
    """ 下载数据 """
    data_download = driver.find_element_by_xpath('//*[@id="id_index_view"]/div[1]/div/div/div[2]/div[5]/button[2]')
    # 在按钮上有东西挡住了，需要用JS执行点击
    driver.execute_script("$(arguments[0]).click()", data_download) 
    time.sleep(10)

    # data_download = 'document.querySelector("#id_index_view > div.main.ng-scope > div > div > div:nth-child(2) > div.btn-group-info.col-xs-7 > button:nth-child(2)").click()'
    # driver.execute_script(data_download)
    # time.sleep(10)

    ribao_path = "E://tencent_ieg_work//王者荣耀//设备和用户画像//6.代练//DailianModel//result"
    ribao_file = ""

    file_list = os.listdir(ribao_path)
    for file_name in file_list:
        if "日报同步" in file_name and sub_time in file_name:
            ribao_file = os.path.join(ribao_path, file_name)

    # "5人举报"
    cur_jubao_5ren, rate_jubao_5ren = cal_rate(ribao_file, sheet_nums=0, col_nums="5人举报活跃占比（百万分）")

    # Manifest
    cur_manifest, rate_manifest = cal_rate(ribao_file, sheet_nums=2, col_nums="预测为黑值的Manifest")
    cur_manifest_guize, _ = cal_rate(ribao_file, sheet_nums=4, col_nums="覆盖率")

    # 破解版
    _, rate_anzhuo_pojieban = cal_rate(ribao_file, sheet_nums=6, col_nums="破解版使用总量")
    _, rate_iso_pojieban = cal_rate(ribao_file, sheet_nums=7, col_nums="ios破解版使用总量")

    # replay
    cur_replay, _ = cal_rate(ribao_file, sheet_nums=8, col_nums="处罚对局数")
    cur_5_replay, _ = cal_rate(ribao_file, sheet_nums=9, col_nums="top10英雄5+举报处罚人数占比")

    # 高端局举报
    _, rate_gaoduan_1 = cal_rate(ribao_file, sheet_nums=12, col_nums="单局被举报1+占比（万分之）")
    _, rate_gaoduan_2 = cal_rate(ribao_file, sheet_nums=13, col_nums="单局被举报2+占比（万分之）")

    # 安卓平台举报
    _, rate_jubao_anzhuo = cal_rate(ribao_file, sheet_nums=14, col_nums="安卓")
    _, rate_jubao_ios = cal_rate(ribao_file, sheet_nums=14, col_nums="iOS")

    # 体验服5+举报
    _, rate_tiyanfu_jubao = cal_rate(ribao_file, sheet_nums=15, col_nums="5人举报")

    """写文档"""
    word_file = "./result/日报.docx"
    document = docx.Document()

    # 往文档中添加段落
    document.add_paragraph("@leanchen(陈旺林)  @akali(李冠龙)  @panpan(潘旭)  @edwarddeng(邓立丰)  @aylinyang(杨琳淇)  @eijiexiao(肖跃坚)  @zhiyuanlu(鲁芝渊)  ")
    document.add_paragraph("中心KPI达标 (演员专项+外挂对抗) ")
    document.add_paragraph('1.外挂专项：')
    if rate_jubao_5ren > 0:
        document.add_paragraph('（1）外挂5人举报活跃占比上涨，占比百万分之%.2f，上涨%.1f%s↑：' % (cur_jubao_5ren, abs(rate_jubao_5ren * 100), "%"))
    else:
        document.add_paragraph('（1）外挂5人举报活跃占比下降，占比百万分之%.2f，下降%.1f%s↓：' % (cur_jubao_5ren, abs(rate_jubao_5ren * 100), "%"))
    document.add_picture('./result/test_V11.png', width = Inches(7))

    rate_manifest = write_content_in_doc(rate_manifest)
    document.add_paragraph('（2）manifest可疑监控%s，离线规则覆盖提单的%.1f%s：' % (rate_manifest, cur_manifest_guize * 100, "%"))
    document.add_picture('./result/test_V22.png', width = Inches(7))
    document.add_picture('./result/test_V23.png', width = Inches(7))

    rate_anzhuo_pojieban = write_content_in_doc(rate_anzhuo_pojieban)
    rate_iso_pojieban = write_content_in_doc(rate_iso_pojieban)
    document.add_paragraph('（3）破解版监控使用量%s，ios破解版监控使用量%s' % (rate_anzhuo_pojieban, rate_iso_pojieban))
    document.add_picture('./result/test_V33.png', width = Inches(7))

    document.add_paragraph('（4）replay外挂检测，昨天replay处罚人局总量为%d，占5+举报%.1f%s：' % (cur_replay, cur_5_replay * 100, "%"))
    document.add_picture('./result/test_V44.png', width = Inches(7))

    rate_gaoduan_1 = write_content_in_doc(rate_gaoduan_1)
    rate_gaoduan_2 = write_content_in_doc(rate_gaoduan_2)
    document.add_paragraph('（5）高端局中，单局举报1+占比%s，2+举报%s：' % (rate_gaoduan_1, rate_gaoduan_2))
    document.add_picture('./result/test_V55.png', width = Inches(7))

    rate_jubao_anzhuo = write_content_in_doc(rate_jubao_anzhuo)
    rate_jubao_ios = write_content_in_doc(rate_jubao_ios)
    document.add_paragraph('（6）五人举报中，安卓平台%s, IOS平台%s：' % (rate_jubao_anzhuo, rate_jubao_ios))
    document.add_picture('./result/test_V66.png', width = Inches(7))

    document.add_paragraph('（8）高举报APK监控，top3高可疑为：com.ziggurat.xdsmoba5v5、com.jingxintech.jxz、com.android.launcher2')
    
    rate_tiyanfu_jubao = write_content_in_doc(rate_tiyanfu_jubao)
    document.add_paragraph('（9）体验服5+举报人数%s' % (rate_tiyanfu_jubao))
    document.add_picture('./result/test_V77.png', width = Inches(7))
    
    document.save(word_file)

    print("Finish!!!!!!!!!!!!1")

    # 删除下载下来的数据
    result = os.popen("rm -r %s" % (ribao_file))
    print("Delete Data File")
